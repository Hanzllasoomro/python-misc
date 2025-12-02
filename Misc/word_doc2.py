from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document for the visual table format based on the user-provided structure
doc = Document()
doc.add_heading('Hotel Chain Data Warehouse Solution - Information Package Diagram', 0)

# Adding a description paragraph
doc.add_paragraph("The following table represents the Information Package Diagram for the hotel chain data warehouse solution, "
                  "showing dimensions, hierarchies, and facts for analyzing hotel occupancy patterns.")

# Create the main table similar to the provided format
table = doc.add_table(rows=13, cols=5)
table.style = 'Table Grid'

# Set column headers for dimensions
headers = ["Time", "Check-in / Check-out", "Room", "Room Occupancy", "Rate Plan / Travel Agent"]
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Fill in dimension data with hierarchies/categories as per user structure
time_data = ["Year", "Quarter", "Month", "Date", "Day of Week", "Holiday Indicator"]
check_in_out_data = ["Check-in", "Check-out", "Occupancy"]
room_data = ["Room Type", "Room ID", "Max Occupancy"]
room_occupancy_data = ["Occupancy Count", "Room Revenue", "Occupancy Rate", "Date (FK)", "Room ID (FK)", "Customer ID (FK)"]
rate_plan_travel_agent_data = ["Package ID", "Package Description", "Package Type", "Seasonal Pricing", "Agent ID", "Agency Name", "Commission Rate"]

# Fill table cells with dimension data
dimension_data = [time_data, check_in_out_data, room_data, room_occupancy_data, rate_plan_travel_agent_data]
for col, data in enumerate(dimension_data):
    for row, item in enumerate(data, start=1):
        table.cell(row, col).text = item

# Merge cells for facts section
fact_cell = table.cell(12, 0)
fact_cell.merge(table.cell(12, 4))
fact_cell.text = "Facts: Occupancy Count, Room Revenue, Occupancy Rate"
fact_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
fact_cell.paragraphs[0].runs[0].bold = True
fact_cell.paragraphs[0].runs[0].font.size = Pt(12)

# Save the document
file_path_table_format = "Hotel_Chain_DWH_Information_Package_Table.docx"  # Save in the current directory
doc.save(file_path_table_format)

file_path_table_format
