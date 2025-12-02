from docx import Document

# Create a new document
doc = Document()

doc.add_heading('University Exam Questions and Answers', level=1)

# Question 1 (a)
doc.add_heading('Question 1 (a):', level=2)
doc.add_paragraph("""
Identify the QA along with its six parts for the following general scenario in concern:
“A user, wanting to minimize the impact of an error, wishes to cancel a system operation at runtime; cancellation takes place in less than one second”.
""")

doc.add_heading('Answer:', level=2)
doc.add_paragraph("""
**Quality Attribute (QA):** Usability

**Six Parts for QA:**
1. **Stimulus**: A user wishes to cancel a system operation.
2. **Response**: The cancellation operation must be performed.
3. **Environment**: Runtime.
4. **Artifact**: The user interface and system.
5. **Response Measure**: The operation should be canceled in less than one second.
6. **Stakeholders**: End-users and system designers.
""")

# Question 1 (b)
doc.add_heading('Question 1 (b):', level=2)
doc.add_paragraph("""
Summarize the concept of architectural tactics in software architecture and design. Briefly describe any 3 tactics of your choice for making your system highly available and/or secure.
""")

doc.add_heading('Answer:', level=2)
doc.add_paragraph("""
**Architectural Tactics**:
Architectural tactics are design decisions used to achieve specific quality attribute goals in software systems. They help in addressing concerns like availability, security, performance, and usability.

**Three Tactics**:
1. **Redundancy (High Availability)**:
   - Duplicate critical system components to ensure continuity in case of failures.
   - Example: Using backup servers.

2. **Authentication and Authorization (Security)**:
   - Enforce user identification and access control mechanisms.
   - Example: Role-based access control systems.

3. **Input Validation (Security)**:
   - Validate user inputs to prevent malicious activities such as SQL injection.
   - Example: Implementing regex-based validation.
""")

# Question 2 (a)
doc.add_heading('Question 2 (a):', level=2)
doc.add_paragraph("""
Discuss the idea of software design patterns along with its categories. Also list the design patterns in each category.
""")

doc.add_heading('Answer:', level=2)
doc.add_paragraph("""
**Software Design Patterns**:
Design patterns provide reusable solutions for common software design problems.

**Categories and Patterns**:
1. **Creational Patterns**:
   - Examples: Singleton, Factory Method, Builder, Prototype, Abstract Factory.

2. **Structural Patterns**:
   - Examples: Adapter, Bridge, Composite, Decorator, Facade, Flyweight, Proxy.

3. **Behavioral Patterns**:
   - Examples: Chain of Responsibility, Command, Interpreter, Iterator, Mediator, Memento, Observer, State, Strategy, Template Method, Visitor.
""")

# Question 2 (b)
doc.add_heading('Question 2 (b):', level=2)
doc.add_paragraph("""
Develop code for any sample application utilizing the Singleton design pattern.
""")

doc.add_heading('Answer:', level=2)
doc.add_paragraph("""
**Singleton Design Pattern Code Example in Java**:

```java
public class Singleton {
    // Static instance of the Singleton class
    private static Singleton instance;

    // Private constructor to prevent instantiation
    private Singleton() {}

    // Public method to provide access to the instance
    public static Singleton getInstance() {
        if (instance == null) {
            instance = new Singleton();
        }
        return instance;
    }

    public void showMessage() {
        System.out.println("Singleton Pattern Example!");
    }
}

// Usage
public class Main {
    public static void main(String[] args) {
        Singleton singleton = Singleton.getInstance();
        singleton.showMessage();
    }
}
""")


doc.save("e:\python\Dimensional_Modeling_Assignment_Descriptive.docx")