from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('Comparison of Testing Tools: Selenium, Appium, JUnit, TestNG, and Cypress', level=1)

# Add a description
doc.add_paragraph(
    "Below is a detailed comparison of five popular testing tools: Selenium, Appium, JUnit, TestNG, and Cypress. "
    "This table includes detailed features, use cases, strengths, and limitations to help you understand their differences and applications better."
)

# Add a table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Add table headers
headers = table.rows[0].cells
headers[0].text = 'Feature/Tool'
headers[1].text = 'Selenium'
headers[2].text = 'Appium'
headers[3].text = 'JUnit'
headers[4].text = 'TestNG'
headers[5].text = 'Cypress'

# Add table rows
rows = [
    ("Primary Use Case", "Automated testing of web applications across multiple browsers and platforms.",
     "Automated testing of mobile applications (native, hybrid, and mobile web) on iOS and Android.",
     "Unit testing for Java applications, primarily used by developers for testing small code units.",
     "Advanced testing framework for unit, integration, and end-to-end testing in Java.",
     "End-to-end testing for modern web applications, focusing on simplicity and speed."),
    ("Language Support", "Supports multiple languages: Java, Python, C#, Ruby, JavaScript, and more.",
     "Supports multiple languages: Java, Python, C#, Ruby, JavaScript, and more.",
     "Primarily designed for Java.", "Primarily designed for Java.", "Exclusively supports JavaScript/TypeScript."),
    ("Platform Support", "Cross-platform: Windows, macOS, Linux.", "Cross-platform: Supports iOS, Android, and Windows apps.",
     "Cross-platform: Works on any platform that supports Java.", "Cross-platform: Works on any platform that supports Java.",
     "Cross-platform: Windows, macOS, Linux."),
    ("Open Source", "Yes, open-source and free to use.", "Yes, open-source and free to use.", "Yes, open-source and free to use.",
     "Yes, open-source and free to use.", "Yes, open-source and free to use."),
    ("Parallel Execution", "Supports parallel execution using Selenium Grid.", "Supports parallel execution using Appium Grid.",
     "Limited support for parallel execution (requires additional setup).", "Built-in support for parallel execution of tests.",
     "Built-in support for parallel execution."),
    ("Integration", "Integrates with CI/CD tools (Jenkins, GitLab), TestNG, JUnit, and reporting tools (ExtentReports).",
     "Integrates with CI/CD tools, TestNG, JUnit, and reporting tools.", "Integrates with build tools like Maven, Gradle, and CI/CD pipelines.",
     "Integrates with Maven, Gradle, CI/CD tools, and reporting frameworks.", "Integrates with CI/CD tools (Jenkins, GitHub Actions) and supports cloud testing platforms."),
    ("Ease of Use", "Moderate learning curve due to the need for programming skills and setup.",
     "Moderate learning curve, especially for mobile-specific configurations.", "Easy to use for developers familiar with Java.",
     "Moderate learning curve, but easier than Selenium for advanced features.", "Easy to use with a developer-friendly API and real-time reloading."),
    ("Community Support", "Large and active community with extensive documentation and forums.",
     "Large and active community, especially for mobile testing.", "Large and active community, widely used in Java development.",
     "Large and active community, widely used in Java testing.", "Growing community with strong support from developers and testers."),
    ("Reporting", "Basic reporting; requires third-party tools (e.g., ExtentReports, Allure) for advanced reports.",
     "Basic reporting; requires third-party tools for advanced reports.", "Basic reporting; often extended with plugins or integrations.",
     "Advanced built-in reporting with HTML reports and custom listeners.", "Advanced built-in reporting with screenshots, videos, and real-time test execution logs."),
    ("Browser Support", "Supports all major browsers: Chrome, Firefox, Safari, Edge, etc.", "N/A (focused on mobile platforms).",
     "N/A (unit testing framework).", "N/A (unit and integration testing framework).", "Supports Chrome, Firefox, Edge, and Electron."),
    ("Mobile Testing", "No native support for mobile testing.", "Yes, supports native, hybrid, and mobile web applications.",
     "No native support for mobile testing.", "No native support for mobile testing.", "No native support for mobile testing."),
    ("Strengths", "- Cross-browser compatibility.\n- Supports multiple programming languages.\n- Highly customizable.",
     "- Cross-platform mobile testing.\n- Supports multiple languages.\n- Integrates with Selenium.",
     "- Simple and lightweight.\n- Ideal for unit testing.\n- Strong integration with Java tools.",
     "- Advanced features like parallel execution, grouping, and prioritization.\n- Flexible annotations.",
     "- Fast execution.\n- Real-time reloading.\n- Excellent debugging capabilities.\n- Easy setup."),
    ("Limitations", "- Steeper learning curve.\n- Requires additional tools for reporting and parallel execution.",
     "- Complex setup for mobile environments.\n- Slower execution compared to web testing tools.",
     "- Limited to unit testing.\n- Lacks advanced features for end-to-end testing.", "- Requires Java knowledge.\n- Can be overwhelming for beginners.",
     "- Limited browser support (no Safari or IE).\n- Only supports JavaScript/TypeScript.")
]

# Add rows to the table
for row in rows:
    row_cells = table.add_row().cells
    for i, content in enumerate(row):
        row_cells[i].text = content

# Add references
doc.add_heading('References', level=2)
doc.add_paragraph(
    "1. Selenium: [Selenium Official Documentation](https://www.selenium.dev/documentation/)\n"
    "2. Appium: [Appium Official Documentation](http://appium.io/docs/en/about-appium/intro/)\n"
    "3. JUnit: [JUnit 5 User Guide](https://junit.org/junit5/docs/current/user-guide/)\n"
    "4. TestNG: [TestNG Documentation](https://testng.org/doc/documentation-main.html)\n"
    "5. Cypress: [Cypress Documentation](https://docs.cypress.io/guides/overview/why-cypress)"
)

# Save the document
doc.save('Testing_Tools_Comparison.docx')

print("Word document generated successfully!")