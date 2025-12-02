from docx import Document

# Create a new Document
doc = Document()
doc.add_heading("SDA PBL Assignment", 0)

# Question 1 - Quality Attribute Analysis in a Software System
doc.add_heading("Question 1: Quality Attribute Analysis in a Software System", level=1)

# Description of the system selected
doc.add_paragraph("System: E-commerce Platform\n"
                  "The chosen system is an e-commerce platform, which allows users to browse products, make purchases, "
                  "and manage orders. This platform handles high traffic and sensitive user data, making quality attributes "
                  "like reliability, security, usability, and performance essential.")

# Quality Attribute Analysis
# Reliability
doc.add_heading("Reliability", level=2)
doc.add_paragraph("Reliability is crucial for the e-commerce platform to ensure users can access the site at all times, "
                  "particularly during peak shopping periods (e.g., Black Friday).\n\n"
                  "System-Specific Scenario: A user tries to make a purchase during a flash sale event. The system must "
                  "handle increased traffic without downtime or crashing.\n\n"
                  "General Scenario: For any software system, reliability means the system consistently performs its "
                  "required functions under both normal and extreme conditions.\n")

# Security
doc.add_heading("Security", level=2)
doc.add_paragraph("Security is critical to protect users' personal and payment information.\n\n"
                  "System-Specific Scenario: A user enters sensitive information such as credit card details. The system "
                  "ensures data encryption and protects against unauthorized access.\n\n"
                  "General Scenario: Security measures, such as encryption and authentication, protect data in any "
                  "software system, reducing the risk of unauthorized access.\n")

# Usability
doc.add_heading("Usability", level=2)
doc.add_paragraph("Usability impacts user satisfaction and platform navigation ease.\n\n"
                  "System-Specific Scenario: A user attempts to locate and purchase a product quickly. An intuitive layout "
                  "and clear navigation paths are critical.\n\n"
                  "General Scenario: Usability in any system makes it easier for users to perform tasks without requiring "
                  "additional assistance, enhancing user experience.\n")

# Performance
doc.add_heading("Performance", level=2)
doc.add_paragraph("Performance affects system speed, especially with large volumes of users.\n\n"
                  "System-Specific Scenario: During a high-traffic period, the system responds to user requests without lag, "
                  "ensuring a smooth shopping experience.\n\n"
                  "General Scenario: In any system, performance metrics like response time and processing speed are key "
                  "to ensuring efficient operations under various load conditions.\n")

# Question 2 - Scenario Development for an Additional Quality Attribute
doc.add_heading("Question 2: Scenario Development for an Additional Quality Attribute - Scalability", level=1)

# Scalability
doc.add_paragraph("Scalability is the software's ability to handle growing user demand or data volume without performance degradation.\n")

# General Scenario
doc.add_heading("General Scenario", level=2)
doc.add_paragraph("Source: User\n"
                  "Stimulus: An increase in user activity (e.g., during a sales event).\n"
                  "Environment: Online platform.\n"
                  "Artifact: Server infrastructure and database handling user requests.\n"
                  "Response: The system scales up resources to accommodate increased traffic.\n"
                  "Response Measure: System maintains acceptable response times under increased load.\n")

# Concrete Scenario
doc.add_heading("Concrete Scenario", level=2)
doc.add_paragraph("In an e-commerce platform, scalability allows the system to support an influx of users during peak "
                  "times, like holiday sales, by dynamically increasing server capacity.\n"
                  "Source: User traffic.\n"
                  "Stimulus: Surge in transactions and concurrent users.\n"
                  "Environment: E-commerce platform during a flash sale.\n"
                  "Artifact: Backend servers and database.\n"
                  "Response: System deploys additional resources to handle the load.\n"
                  "Response Measure: System maintains response times below 2 seconds for all user actions.\n")

# Save the document
doc_path = r"E:\SDA_PBL_Assignment.docx"
doc.save(doc_path)

doc_path
