from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_table(doc, title, headers, rows):
    doc.add_paragraph(title).bold = True
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = True

    # Set header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, item in enumerate(row_data):
            table.cell(row_idx, col_idx).text = item
            table.cell(row_idx, col_idx).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return doc

# Create a new Word document
doc = Document()

# Add heading and introduction
doc.add_heading('Data Warehouse Solution for Hotel Chain - Visual Layouts', 0)
doc.add_paragraph("This document provides a visual representation of the Information Package Diagram for the hotel chain data warehouse solution. It includes tables to represent each dimension and fact table relationship for a clearer understanding of the schema.")

# Add a heading for the Information Package Diagram
doc.add_heading('Information Package Diagram', level=1)
doc.add_paragraph("Below is the visual representation of the Information Package Diagram using tables:")

# Create tables for each dimension and fact table
create_table(doc, "Time Dimension", ["Attributes"], [
    ["Date"], ["Day of Week"], ["Month"], ["Quarter"], ["Year"], ["Holiday Indicator"]
])

# ... (similarly for other dimensions and the fact table)

# Save the document
file_path_visual_tables = "Hotel_Chain_DWH_Solution_Visual_Tables.docx"
doc.save(file_path_visual_tables)